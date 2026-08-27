# -*- coding: utf-8 -*-
"""
產出「八.近兩年相關新聞」Word 檔
=================================

完全比照 115_A-I-001_54955208_台灣禾邦電子有限公司.doc 的版面：

    八.近兩年相關新聞           <- 樣式「廠商 - 目錄 標 1」，18pt 粗體
    新聞目錄                    <- 置中
    {TOC 目錄欄位，含頁碼}       <- TOC \\h \\z \\t "title2.0,1"
    (空行)
    新聞標題                    <- 樣式「title2.0」，16pt 粗體（目錄由此產生）
    來源行                      <- 樣式「廠商 - 新聞來源」，11pt
    副標／小標                  <- 樣式「廠商 - 新聞 2」，13pt 粗體
    內文段落                    <- 樣式「廠商 - 新聞內容」，13pt、固定行高 24pt、首行縮排 2 字元
    ...

樣式規格取自原始 .doc（Word COM 讀出），在此以程式重建，不需範本檔，
Linux/Streamlit Cloud 也能執行。
"""

from __future__ import annotations

__version__ = "1.5.3"

import os
import re
import subprocess
import sys
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# 樣式名稱（與原始 .doc 一致）
S_HEAD = "廠商 - 目錄 標 1"
S_TITLE = "title2.0"
S_SOURCE = "廠商 - 新聞來源"
S_BODY = "廠商 - 新聞內容"
S_SUB = "廠商 - 新聞 2"

ASCII_FONT = "Times New Roman"
CJK_BODY = "標楷體"
CJK_TITLE = "Microsoft YaHei"


# --------------------------------------------------------------------------- #
# 樣式建置
# --------------------------------------------------------------------------- #
def _set_fonts(style, ascii_font: str, cjk_font: str) -> None:
    style.font.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)


def _first_line_chars(paragraph_format, chars: int) -> None:
    """設定「首行縮排 N 字元」(w:firstLineChars)，python-docx 未直接支援。"""
    ppr = paragraph_format.element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))


def _mk_style(doc, name: str, *, size: float, bold: bool, cjk: str,
              align, spacing_rule, spacing_pt: Optional[float],
              before: float = 0, after: float = 0,
              first_line_chars: int = 0):
    styles = doc.styles
    try:
        st = styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles["Normal"]
    st.font.size = Pt(size)
    st.font.bold = bold
    _set_fonts(st, ASCII_FONT, cjk)
    pf = st.paragraph_format
    pf.alignment = align
    if spacing_pt is not None:
        pf.line_spacing = Pt(spacing_pt)
    pf.line_spacing_rule = spacing_rule
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line_chars:
        _first_line_chars(pf, first_line_chars)
    st.quick_style = True
    return st


def _build_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    _set_fonts(normal, ASCII_FONT, CJK_BODY)

    _mk_style(doc, S_HEAD, size=18, bold=True, cjk=CJK_BODY,
              align=WD_ALIGN_PARAGRAPH.LEFT,
              spacing_rule=WD_LINE_SPACING.AT_LEAST, spacing_pt=20,
              before=2.5, after=2.5)
    _mk_style(doc, S_TITLE, size=16, bold=True, cjk=CJK_TITLE,
              align=WD_ALIGN_PARAGRAPH.LEFT,
              spacing_rule=WD_LINE_SPACING.AT_LEAST, spacing_pt=12,
              before=12, after=0)
    _mk_style(doc, S_SOURCE, size=11, bold=False, cjk=CJK_BODY,
              align=WD_ALIGN_PARAGRAPH.LEFT,
              spacing_rule=WD_LINE_SPACING.AT_LEAST, spacing_pt=12)
    _mk_style(doc, S_BODY, size=13, bold=False, cjk=CJK_BODY,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              spacing_rule=WD_LINE_SPACING.EXACTLY, spacing_pt=24,
              first_line_chars=2)
    _mk_style(doc, S_SUB, size=13, bold=True, cjk=CJK_BODY,
              align=WD_ALIGN_PARAGRAPH.LEFT,
              spacing_rule=WD_LINE_SPACING.SINGLE, spacing_pt=None,
              before=2.5, after=2.5)


def _page_setup(doc) -> None:
    sec = doc.sections[0]
    sec.page_width = Pt(595.3)      # A4
    sec.page_height = Pt(841.9)
    sec.left_margin = Pt(90)
    sec.right_margin = Pt(90)
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)


def _enable_update_fields(doc) -> None:
    """讓 Word 開檔時自動更新目錄頁碼。"""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def _add_toc_field(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\h \\z \\t "%s,1" ' % S_TITLE
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "【請在 Word 中按 Ctrl+A 後 F9 更新目錄頁碼】"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def _add_page_numbers(doc) -> None:
    """頁尾置中頁碼。"""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)


# --------------------------------------------------------------------------- #
# 寫入前的字元清理
# --------------------------------------------------------------------------- #
# XML 1.0 只允許 \t \n \r 這三個控制字元，其餘 C0 控制碼、￾/￿ 與落單的
# surrogate 一律不合法。新聞網頁偶爾會夾帶這些字元（多半來自編輯器或轉碼殘留），
# 肉眼看不出來，但 lxml 在寫入時會丟 ValueError／UnicodeEncodeError，導致整份
# Word 檔產不出來。既然它們在 Word 裡本來也顯示不了，寫入前直接濾掉。
_ILLEGAL_XML = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff\ufffe\uffff]")


def xml_safe(text) -> str:
    """濾掉不能寫進 XML 的字元，並把 \\r\\n／\\r 正規化成 \\n。"""
    if not text:
        return ""
    s = str(text).replace("\r\n", "\n").replace("\r", "\n")
    return _ILLEGAL_XML.sub("", s)


# --------------------------------------------------------------------------- #
# 來源行
# --------------------------------------------------------------------------- #
def source_lines(article, style: str = "auto") -> List[str]:
    """組出新聞來源標示行。

    print 體例：【2025-07-31/經濟日報/A15版/經營管理】【尹慧中】
    web   體例：網址 + 「MoneyDJ新聞 2025-06-18 11:23:41 記者 萬惠雯 報導」
    """
    date = article.date_str
    src = article.source or ""
    author = (article.author or "").strip()

    if style == "print" or (style == "auto" and not article.url):
        head = "【%s】" % "/".join(x for x in [date, src] if x)
        return [head + ("【%s】" % author if author else "")]

    line2_bits = [x for x in [src, date] if x]
    if author:
        line2_bits.append("記者%s 報導" % author if not author.startswith("記者") else author)
    lines = []
    if article.url:
        lines.append(article.url)
    if line2_bits:
        lines.append(" ".join(line2_bits))
    return lines or ["【%s】" % date]


# --------------------------------------------------------------------------- #
# 主要建置函式
# --------------------------------------------------------------------------- #
def build_docx(articles: Iterable,
               out_path: str,
               section_title: str = "八.近兩年相關新聞",
               catalog_title: str = "新聞目錄",
               source_style: str = "auto",
               company: str = "") -> str:
    """把 Article 清單寫成 Word 檔（.docx），回傳輸出路徑。"""
    doc = Document()
    _page_setup(doc)
    _build_styles(doc)
    _enable_update_fields(doc)
    _add_page_numbers(doc)

    doc.add_paragraph(xml_safe(section_title), style=S_HEAD)

    p = doc.add_paragraph(xml_safe(catalog_title))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_toc_field(doc)

    for art in articles:
        doc.add_paragraph(xml_safe(art.title), style=S_TITLE)
        for line in source_lines(art, source_style):
            doc.add_paragraph(xml_safe(line), style=S_SOURCE)
        if getattr(art, "subtitle", ""):
            doc.add_paragraph(xml_safe(art.subtitle), style=S_SUB)
        for para in art.paragraphs:
            text = xml_safe(para).strip()
            if not text:
                continue
            # 短句且無句號者視為小標
            if len(text) <= 22 and not re.search(r"[。，；：？！]", text):
                doc.add_paragraph(text, style=S_SUB)
            else:
                doc.add_paragraph(text, style=S_BODY)

    out_path = os.path.abspath(out_path)
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


# --------------------------------------------------------------------------- #
# 檔名與 .doc 轉檔
# --------------------------------------------------------------------------- #
def make_filename(year: str, case_no: str, tax_id: str, company: str,
                  ext: str = ".doc", sep: str = "_") -> str:
    """組出 115_A-I-001_54955208_台灣禾邦電子有限公司.doc 這種檔名。"""
    parts = [xml_safe(x).strip() for x in (year, case_no, tax_id, company)
             if xml_safe(x).strip()]
    name = sep.join(parts)
    name = re.sub(r'[\\/:*?"<>|]', "", name)
    return name + ext


def convert_to_doc(docx_path: str, doc_path: Optional[str] = None) -> Optional[str]:
    """在有安裝 Word 的 Windows 上把 .docx 另存為 .doc（wdFormatDocument97=0），
    同時更新目錄頁碼。沒有 Word（例如 Streamlit Cloud）時回傳 None。"""
    if sys.platform != "win32":
        return None
    docx_path = os.path.abspath(docx_path)
    doc_path = os.path.abspath(doc_path or os.path.splitext(docx_path)[0] + ".doc")
    ps = r"""
$ErrorActionPreference = 'Stop'
$w = New-Object -ComObject Word.Application
$w.Visible = $false
$w.DisplayAlerts = 0
try {{
  $d = $w.Documents.Open("{src}", $false, $false)
  foreach ($f in $d.Fields) {{ $null = $f.Update() }}
  $d.TablesOfContents | ForEach-Object {{ $_.Update() }}
  $d.SaveAs([ref]"{dst}", [ref]0)
  $d.Close(0)
}} finally {{
  $w.Quit()
}}
""".format(src=docx_path.replace("\\", "\\\\"), dst=doc_path.replace("\\", "\\\\"))
    try:
        subprocess.run(["powershell", "-NoProfile", "-NonInteractive", "-Command", ps],
                       check=True, capture_output=True, timeout=180)
    except Exception:      # noqa: BLE001
        return None
    return doc_path if os.path.exists(doc_path) else None
